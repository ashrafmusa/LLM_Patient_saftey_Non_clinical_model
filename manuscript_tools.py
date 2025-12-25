import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
import re

def analyze_manuscript(file_path):
    """Analyze uploaded manuscript document"""
    try:
        doc = Document(file_path)
        
        # Extract text
        full_text = "\n".join([p.text for p in doc.paragraphs])
        word_count = len(full_text.split())
        
        # Extract tables
        table_count = len(doc.tables)
        
        # Extract references (simple heuristic - numbered list at end)
        ref_pattern = r'^\d+\.\s'
        references = len([p for p in doc.paragraphs if re.match(ref_pattern, p.text)])
        
        # Extract title (first paragraph usually)
        title = doc.paragraphs[0].text if doc.paragraphs else "N/A"
        
        return {
            "title": title[:100],
            "word_count": word_count,
            "tables": table_count,
            "estimated_references": references,
            "paragraphs": len(doc.paragraphs),
            "font_sizes": extract_font_info(doc)
        }
    except Exception as e:
        return {"error": str(e)}

def extract_font_info(doc):
    """Extract font information from document"""
    fonts = {}
    for para in doc.paragraphs:
        for run in para.runs:
            font_name = run.font.name or "Default"
            size = run.font.size
            if size:
                size = int(size.pt)
            fonts[font_name] = size
    return fonts

def check_compliance(manuscript_data, journal_requirements):
    """Check manuscript against journal requirements"""
    issues = []
    warnings = []
    
    # Word count check
    if manuscript_data["word_count"] > journal_requirements["word_limit"]:
        issues.append(f"❌ Word count {manuscript_data['word_count']} exceeds limit of {journal_requirements['word_limit']}")
    elif manuscript_data["word_count"] < journal_requirements["word_limit"] * 0.8:
        warnings.append(f"⚠️  Word count {manuscript_data['word_count']} is significantly below limit")
    else:
        issues.append(f"✅ Word count {manuscript_data['word_count']} is within limits")
    
    # Tables check
    if manuscript_data["tables"] > journal_requirements["tables_limit"]:
        issues.append(f"❌ Tables {manuscript_data['tables']} exceed limit of {journal_requirements['tables_limit']}")
    else:
        issues.append(f"✅ Tables {manuscript_data['tables']} are within limits")
    
    # References check
    if manuscript_data["estimated_references"] > journal_requirements["references_limit"]:
        issues.append(f"❌ References {manuscript_data['estimated_references']} exceed limit of {journal_requirements['references_limit']}")
    else:
        issues.append(f"✅ References {manuscript_data['estimated_references']} are within limits")
    
    return {"issues": issues, "warnings": warnings}

def generate_template(journal_name, author_name, institution):
    """Generate manuscript template for journal"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_paragraph()
    title.paragraph_format.line_spacing = 2.0
    title_run = title.add_run("[Manuscript Title - Replace with your title]")
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    # Author info (placeholder)
    author_para = doc.add_paragraph()
    author_para.paragraph_format.line_spacing = 2.0
    author_para.add_run(f"Author: {author_name}\n{institution}")
    
    # Abstract
    abstract = doc.add_paragraph("\nAbstract\n")
    abstract.paragraph_format.line_spacing = 2.0
    abstract_run = abstract.runs[0]
    abstract_run.bold = True
    doc.add_paragraph("[Background]\n[Methods]\n[Results]\n[Conclusions]", style='List Number')
    
    # Introduction
    doc.add_paragraph("\nIntroduction\n").runs[0].bold = True
    doc.add_paragraph("[Introduction text...]")
    
    # Methods
    doc.add_paragraph("\nMethods\n").runs[0].bold = True
    doc.add_paragraph("[Methods text...]")
    
    # Results
    doc.add_paragraph("\nResults\n").runs[0].bold = True
    doc.add_paragraph("[Results text...]")
    
    # Discussion
    doc.add_paragraph("\nDiscussion\n").runs[0].bold = True
    doc.add_paragraph("[Discussion text...]")
    
    # References
    doc.add_paragraph("\nReferences\n").runs[0].bold = True
    doc.add_paragraph("[Format references according to journal style]")
    
    return doc

def export_checklist(journal_requirements):
    """Generate submission checklist"""
    checklist = f"""
MANUSCRIPT SUBMISSION CHECKLIST
Journal: {journal_requirements.get('name', 'N/A')}

CONTENT REQUIREMENTS:
☐ Word count: Maximum {journal_requirements.get('word_limit', 'N/A')} words
☐ Abstract: Maximum {journal_requirements.get('abstract_limit', 'N/A')} words
☐ References: Maximum {journal_requirements.get('references_limit', 'N/A')}
☐ Tables: Maximum {journal_requirements.get('tables_limit', 'N/A')}
☐ Title: Maximum {journal_requirements.get('title_limit', 'N/A')} characters
☐ Keywords: {journal_requirements.get('keywords', 'N/A')} terms

FORMATTING:
☐ Format: {journal_requirements.get('format', 'N/A')}
☐ Style: {journal_requirements.get('style_guide', 'N/A')}
☐ Reporting Guideline: {journal_requirements.get('reporting_guideline', 'N/A')}

AUTHOR REQUIREMENTS:
☐ Author affiliations provided
☐ Corresponding author contact information
☐ Conflict of interest statement
☐ Funding disclosure
☐ CRediT roles (if applicable)

FILES TO SUBMIT:
☐ Main manuscript
☐ Title page (with author info)
☐ Abstract
☐ References
☐ Tables and figures
☐ Supplementary materials
☐ Cover letter
☐ Conflict of interest form
☐ Compliance checklist (STROBE/CONSORT if applicable)

SUBMISSION PORTAL:
{journal_requirements.get('url', 'Check journal website')}
    """
    return checklist
